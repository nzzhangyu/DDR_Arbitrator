from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path("artifacts/resume")
DOCX_PATH = OUT_DIR / "张宇-个人简历-优化版.docx"

FONT = "Microsoft YaHei"
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)


def set_run_font(run, size=None, bold=False, color=DARK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_para_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, text="", size=9.2, bold=False, color=DARK, after=2, before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para_spacing(p, before=before, after=after, line=1.05)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=5, after=2, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True, color=BLUE)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    border.append(bottom)
    p_pr.append(border)
    return p


def add_bullet(doc, text, size=8.9, after=1):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, before=0, after=after, line=1.03)
    p.paragraph_format.left_indent = Cm(0.52)
    p.paragraph_format.first_line_indent = Cm(-0.22)
    run = p.add_run(text)
    set_run_font(run, size=size, color=DARK)
    return p


def add_role_line(doc, left, right=None):
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=1, line=1.0)
    run = p.add_run(left)
    set_run_font(run, size=9.4, bold=True, color=DARK)
    if right:
        run = p.add_run("    " + right)
        set_run_font(run, size=8.8, color=MUTED)
    return p


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = DARK

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    bullet.font.size = Pt(8.9)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=1, line=1.0)
    run = p.add_run("张宇")
    set_run_font(run, size=19, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=4, line=1.0)
    for i, text in enumerate(["FPGA开发工程师 / 数字逻辑设计工程师", "18842345241", "nzzhangyu123@163.com"]):
        if i:
            sep = p.add_run("  |  ")
            set_run_font(sep, size=8.8, color=MUTED)
        run = p.add_run(text)
        set_run_font(run, size=8.8, color=MUTED)

    add_heading(doc, "教育背景")
    add_role_line(doc, "北京理工大学｜电子科学与技术｜硕士", "2021.09 - 2024.06")
    add_role_line(doc, "东北大学｜电子信息工程｜本科", "2016.09 - 2020.06")

    add_heading(doc, "个人摘要")
    add_text(
        doc,
        "硕士学历，2年医疗影像设备 FPGA 开发经验，主要参与 CT 探测器数据采集、旋转通信链路和 DDR 缓存调度相关模块开发。熟悉高速数据接收、CDC、FIFO、DSP 流水线、Aurora 链路、DDR 读写仲裁及板级调试流程，具备从 RTL 设计、仿真验证到板级调试和系统联调的问题闭环经验。",
        size=8.9,
        after=2,
    )

    add_heading(doc, "工作经历")
    add_role_line(doc, "东软医疗系统股份有限公司｜电子工程师（FPGA方向）｜CT产品部", "2024.07 - 至今")
    add_bullet(doc, "负责 CT 产品部件 FPGA 代码设计、模块仿真、板级调试与系统联调。", after=0.5)
    add_bullet(doc, "参与 ASIC 数据采集、像素处理、Aurora 高速链路、DDR 缓存仲裁等核心模块开发。", after=0.5)
    add_bullet(doc, "配合硬件工程师完成单板 bring-up、接口联调、时序问题定位和 ILA 抓波分析。", after=0.5)

    add_heading(doc, "项目经历")
    add_role_line(doc, "数据采集接口（DASI）固件开发")
    add_text(
        doc,
        "负责 CT 探测器 ASIC 数据采集链路相关 FPGA 固件开发，实现多通道 LVDS 数据接收、跨时钟域同步、像素处理、数据重排与成帧输出。",
        size=8.7,
        color=MUTED,
        after=1,
    )
    add_bullet(doc, "设计 ASIC 读出状态机，完成多通道 LVDS 串行数据接收、过采样判决和通道对齐。", after=0.5)
    add_bullet(doc, "设计 CDC 与 FIFO 缓冲结构，解决多 bit 数据跨时钟域传输和多通道同步问题。", after=0.5)
    add_bullet(doc, "搭建 DSP 流水线与双口 RAM 结构，实现多帧权重累加、增益校正、暗电流扣除和像素值计算。", after=0.5)
    add_bullet(doc, "基于 ROM 查找表完成空间映射和数据重排，采用乒乓 RAM 实现连续读写与成帧输出。", after=0.5)
    add_text(doc, "技术栈：Verilog/SystemVerilog、LVDS、CDC、FIFO、DSP流水线、双口RAM、乒乓缓存", size=8.6, color=MUTED, after=2)

    add_role_line(doc, "旋转通信板（Rcomm）模块固件开发")
    add_text(
        doc,
        "负责 CT 旋转端数据链路部分固件开发与重构，覆盖数据接收、组帧、缓存、无损传输、异常恢复和 DDR 缓存仲裁。",
        size=8.7,
        color=MUTED,
        after=1,
    )
    add_bullet(doc, "重构 Aurora 8B/10B 链路层接收逻辑，使用弹性 FIFO + 状态机替代原硬延迟线结构，减少无效空拍。", after=0.5)
    add_bullet(doc, "搭建轻量化自仿真环境，注入 skew、bit 翻转、光纤断联等异常场景，验证数据完整性、CRC 报错和链路恢复能力。", after=0.5)
    add_bullet(doc, "设计 DDR 动态水位仲裁机制，实现写侧高水位优先排空、读侧低水位预取和读写公平仲裁；搭建 DDR fast mock，对 refresh、命令反压、读返回断流、可变读延迟和 pending read 场景进行仿真验证。", after=0.5)
    add_bullet(doc, "参与板级调试，通过 ILA 抓取 FIFO 水位、链路状态和 DDR 读写握手信号，定位缓存边界和链路异常问题。", after=0.5)
    add_text(doc, "技术栈：Aurora 8B/10B、DDR、FIFO、状态机、动态水位仲裁、轻量化mock仿真、ILA", size=8.6, color=MUTED, after=2)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
